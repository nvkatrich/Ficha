"""Context preparation and non-destructive DOCX generation."""
from __future__ import annotations

import re
import shutil
import subprocess
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_BREAK

from .config import DOCUMENT_DIR
from .template_store import copy_template_for_edit

TOKEN_PATTERN = re.compile(r"{{\s*([A-Za-zА-Яа-яЁё0-9_.]+)\s*}}")


class DocumentError(RuntimeError):
    pass


def as_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (list, tuple)):
        return ", ".join(as_text(v) for v in value)
    return str(value)


def money(value: Any, currency: str = "") -> str:
    if value is None or str(value).strip() == "":
        return ""
    try:
        formatted = f"{float(str(value).replace(',', '.')):,.2f}".replace(",", " ").replace(".", ",")
        return f"{formatted} {currency}".strip()
    except (ValueError, TypeError):
        return f"{as_text(value)} {currency}".strip()


def join_name(record: dict[str, Any]) -> str:
    return " ".join(filter(None, [as_text(record.get(key)) for key in ("LAST_NAME", "NAME", "SECOND_NAME")])).strip()


def format_date(value: Any) -> str:
    text = as_text(value)
    if not text:
        return ""
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00")).strftime("%d.%m.%Y")
    except ValueError:
        return text[:10]


def build_context(bundle: dict[str, Any], settings: dict[str, Any], overrides: dict[str, Any] | None = None) -> dict[str, Any]:
    deal = bundle.get("deal", {})
    company = bundle.get("company", {})
    contact = bundle.get("contact", {})
    manager = bundle.get("manager", {})
    products = bundle.get("products", [])
    deal_id = as_text(deal.get("ID"))
    currency = as_text(deal.get("CURRENCY_ID"))
    product_sum = sum(_line_total(row) for row in products)

    context: dict[str, Any] = {
        "document": {
            "number": f"{settings.get('document_prefix', 'KP')}-{deal_id}" if deal_id else settings.get("document_prefix", "KP"),
            "date": datetime.now().strftime("%d.%m.%Y"),
        },
        "deal": {
            "id": deal_id,
            "title": as_text(deal.get("TITLE")),
            "amount": money(deal.get("OPPORTUNITY"), currency),
            "currency": currency,
            "stage": as_text(deal.get("STAGE_ID")),
            "close_date": format_date(deal.get("CLOSEDATE")),
            "comments": as_text(deal.get("COMMENTS")),
        },
        "client": {
            "company": as_text(company.get("TITLE")),
            "inn": as_text(company.get("UF_CRM_")) or as_text(company.get("UF_CRM_INN")),
            "contact_name": join_name(contact),
            "phone": _first_value(contact.get("PHONE")),
            "email": _first_value(contact.get("EMAIL")),
            "address": as_text(company.get("ADDRESS")) or as_text(company.get("ADDRESS_LEGAL")),
        },
        "manager": {
            "name": join_name(manager),
            "phone": as_text(manager.get("PERSONAL_MOBILE")) or as_text(manager.get("WORK_PHONE")),
            "email": as_text(manager.get("EMAIL")),
        },
        "seller": {
            "name": settings.get("company_name", ""),
            "inn": settings.get("company_inn", ""),
            "phone": settings.get("company_phone", ""),
            "email": settings.get("company_email", ""),
            "address": settings.get("company_address", ""),
        },
        "products": [
            {
                "number": index,
                "name": as_text(row.get("PRODUCT_NAME")) or as_text(row.get("PRODUCT_ID")),
                "quantity": as_text(row.get("QUANTITY")),
                "unit": as_text(row.get("MEASURE_NAME")),
                "price": money(row.get("PRICE"), as_text(row.get("CURRENCY")) or currency),
                "total": money(row.get("PRICE_ACCOUNT")) or money(_line_total(row), as_text(row.get("CURRENCY")) or currency),
            }
            for index, row in enumerate(products, 1)
        ],
        "products_total": money(product_sum, currency) if product_sum else money(deal.get("OPPORTUNITY"), currency),
    }
    context["deal"]["custom"] = {str(key): as_text(value) for key, value in deal.items() if str(key).startswith("UF_CRM_")}
    if overrides:
        for key, value in overrides.items():
            if value is not None and str(value).strip():
                set_dotted(context, key, str(value).strip())
    return context


def _line_total(row: dict[str, Any]) -> float:
    try:
        price = row.get("PRICE")
        quantity = row.get("QUANTITY")
        if price not in (None, "") and quantity not in (None, ""):
            return float(str(price).replace(" ", "").replace(",", ".")) * float(str(quantity).replace(" ", "").replace(",", "."))
        return float(str(row.get("PRICE_ACCOUNT", row.get("TOTAL", 0))).replace(" ", "").replace(",", "."))
    except (TypeError, ValueError):
        return 0.0


def _first_value(value: Any) -> str:
    if isinstance(value, list) and value:
        head = value[0]
        return as_text(head.get("VALUE")) if isinstance(head, dict) else as_text(head)
    return as_text(value)


def lookup_dotted(context: dict[str, Any], dotted_key: str) -> str:
    value: Any = context
    for part in dotted_key.split("."):
        if not isinstance(value, dict):
            return ""
        value = value.get(part, "")
    return as_text(value)


def set_dotted(context: dict[str, Any], dotted_key: str, value: Any) -> None:
    parts = dotted_key.split(".")
    target: dict[str, Any] = context
    for part in parts[:-1]:
        child = target.get(part)
        if not isinstance(child, dict):
            child = {}
            target[part] = child
        target = child
    target[parts[-1]] = value


def replace_tokens_in_paragraph(paragraph: Any, context: dict[str, Any]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if not TOKEN_PATTERN.search(full_text):
        return
    replaced = TOKEN_PATTERN.sub(lambda match: lookup_dotted(context, match.group(1)), full_text)
    if paragraph.runs:
        paragraph.runs[0].text = replaced
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(replaced)


def all_paragraphs(container: Any) -> list[Any]:
    output = list(container.paragraphs)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                output.extend(all_paragraphs(cell))
    return output


def replace_tokens(document: Document, context: dict[str, Any]) -> None:
    for paragraph in all_paragraphs(document):
        replace_tokens_in_paragraph(paragraph, context)
    for section in document.sections:
        for header_footer in (section.header, section.footer):
            for paragraph in all_paragraphs(header_footer):
                replace_tokens_in_paragraph(paragraph, context)


def find_products_marker(document: Document) -> Any | None:
    for paragraph in all_paragraphs(document):
        if "{{products.table}}" in paragraph.text:
            return paragraph
    return None


def add_products_table(document: Document, marker: Any, products: list[dict[str, Any]], total: str) -> None:
    marker.text = ""
    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    header = table.rows[0].cells
    labels = ["№", "Наименование", "Кол-во", "Ед.", "Цена", "Сумма"]
    for cell, label in zip(header, labels):
        cell.text = label
    for product in products:
        cells = table.add_row().cells
        values = [product["number"], product["name"], product["quantity"], product["unit"], product["price"], product["total"]]
        for cell, value in zip(cells, values):
            cell.text = as_text(value)
    cells = table.add_row().cells
    cells[0].merge(cells[4]).text = "Итого"
    cells[5].text = total
    # Move appended table directly after the marker in document XML.
    marker._p.addnext(table._tbl)


def generate_docx(template: dict[str, Any], context: dict[str, Any], output_basename: str) -> Path:
    safe_name = re.sub(r"[^A-Za-zА-Яа-яЁё0-9_.-]+", "_", output_basename).strip("._") or "commercial_proposal"
    output_path = DOCUMENT_DIR / f"{safe_name}_{uuid.uuid4().hex[:8]}.docx"
    copy_template_for_edit(template, output_path)
    try:
        document = Document(output_path)
        marker = find_products_marker(document)
        replace_tokens(document, context)
        if marker:
            add_products_table(document, marker, context.get("products", []), context.get("products_total", ""))
        document.save(output_path)
    except Exception as error:
        output_path.unlink(missing_ok=True)
        raise DocumentError("Не удалось собрать документ из этого шаблона.") from error
    return output_path


def convert_to_pdf(docx_path: Path) -> Path:
    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not libreoffice:
        raise DocumentError("PDF-конвертация требует LibreOffice. Установите LibreOffice и повторите попытку.")
    output_dir = docx_path.parent
    result = subprocess.run(
        [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        capture_output=True,
        text=True,
        timeout=90,
        check=False,
    )
    pdf_path = docx_path.with_suffix(".pdf")
    if result.returncode != 0 or not pdf_path.exists():
        raise DocumentError("LibreOffice не смог конвертировать этот файл в PDF.")
    return pdf_path
