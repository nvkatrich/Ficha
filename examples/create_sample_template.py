"""Build a minimal editable example .docx template for the application."""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

output = Path(__file__).with_name("Шаблон_КП_пример.docx")
document = Document()
section = document.sections[0]
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)

style = document.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(10)

document.add_heading("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", level=0)
document.add_paragraph("№ {{document.number}} от {{document.date}}")
document.add_paragraph("Кому: {{client.company}}")
document.add_paragraph("Контактное лицо: {{client.contact_name}}")
document.add_paragraph("\nУважаемые коллеги!\n\nПредлагаем рассмотреть условия по сделке «{{deal.title}}».")
document.add_paragraph("\n{{products.table}}\n")
document.add_paragraph("Общая стоимость: {{products_total}}")
document.add_paragraph("\nС уважением,\n{{manager.name}}\n{{seller.name}}\n{{manager.phone}} · {{manager.email}}")
document.save(output)
print(output)
