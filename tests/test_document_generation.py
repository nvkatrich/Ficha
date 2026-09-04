from pathlib import Path

from docx import Document

from app import document_service as service
from app import template_store as store


def test_generates_document_and_product_table(tmp_path, monkeypatch):
    templates = tmp_path / "templates"
    documents = tmp_path / "documents"
    templates.mkdir()
    documents.mkdir()
    monkeypatch.setattr(store, "TEMPLATE_DIR", templates)
    monkeypatch.setattr(service, "DOCUMENT_DIR", documents)

    source = templates / "test.docx"
    document = Document()
    document.add_paragraph("№ {{document.number}} от {{document.date}}")
    document.add_paragraph("Клиент: {{client.company}}")
    document.add_paragraph("{{products.table}}")
    document.add_paragraph("Итого: {{products_total}}")
    document.save(source)

    context = {
        "document": {"number": "КП-44", "date": "03.09.2026"},
        "client": {"company": "ООО «Клиент»"},
        "products": [{"number": 1, "name": "Консультация", "quantity": "2", "unit": "час", "price": "5 000,00 RUB", "total": "10 000,00 RUB"}],
        "products_total": "10 000,00 RUB",
    }
    output = service.generate_docx({"filename": "test.docx"}, context, "КП клиента")

    assert output.is_file()
    result = Document(output)
    paragraph_text = "\n".join(paragraph.text for paragraph in result.paragraphs)
    assert "КП-44" in paragraph_text
    assert "ООО «Клиент»" in paragraph_text
    assert "{{" not in paragraph_text
    table_text = " ".join(cell.text for table in result.tables for row in table.rows for cell in row.cells)
    assert "Консультация" in table_text
    assert "10 000,00 RUB" in table_text


def test_build_context_accepts_custom_fields_and_overrides():
    bundle = {
        "deal": {"ID": "41", "TITLE": "Внедрение", "OPPORTUNITY": "12000", "CURRENCY_ID": "RUB", "UF_CRM_12": "Особое условие"},
        "company": {"TITLE": "ООО Тест"},
        "contact": {"NAME": "Анна", "LAST_NAME": "Иванова", "PHONE": [{"VALUE": "+79990000000"}]},
        "manager": {"NAME": "Илья", "LAST_NAME": "Петров"},
        "products": [],
    }
    context = service.build_context(bundle, {"document_prefix": "КП"}, {"client.company": "ООО Изменено"})
    assert context["document"]["number"] == "КП-41"
    assert context["deal"]["custom"]["UF_CRM_12"] == "Особое условие"
    assert context["client"]["company"] == "ООО Изменено"
    assert context["client"]["contact_name"] == "Иванова Анна"
